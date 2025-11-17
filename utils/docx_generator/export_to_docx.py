import os
import json
from datetime import datetime
from docx import Document

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

def _pretty_json(obj):
    try:
        return json.dumps(obj, indent=4, ensure_ascii=False)
    except Exception:
        return str(obj)

def _safe_filename(s):
    return "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in s).strip().replace(" ", "_")

def append_test_results_to_docx(tc_no, tc_name, request_details, response_details, assertion_result,
                                output_dir, template_path=None):
    """
    Append a test case report to a docx document and save it in output_dir.
    - template_path: optional path to a .docx template to start from.
    - request_details / response_details are dicts. Request may omit 'Body' for GET.
    """
    # Ensure output dir exists
    os.makedirs(output_dir, exist_ok=True)

    # Load template if exists, otherwise create a new document
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # Add a horizontal separator (heading with timestamp)
    heading = f"Test Case {tc_no} - {tc_name}"
    doc.add_heading(heading, level=1)

    # Request section
    doc.add_heading("Request", level=2)

    # URL
    doc.add_paragraph("\nRequest URL:").runs[0].bold = True
    doc.add_paragraph(str(request_details.get("URL", "")), style="NoSpacing")

    # Method
    doc.add_paragraph("\nRequest Method:").runs[0].bold = True
    doc.add_paragraph(str(request_details.get("Method", "")), style="NoSpacing")

    # Headers
    doc.add_paragraph("\nRequest Headers:").runs[0].bold = True
    headers = request_details.get("Headers", {})
    for k, v in headers.items():
        doc.add_paragraph(f"{k}: {v}", style="NoSpacing")

    # Body (maybe absent for GET)
    if "Body" in request_details and request_details["Body"] is not None:
        doc.add_paragraph("\nBody:").runs[0].bold = True
        doc.add_paragraph(_pretty_json(request_details["Body"]), style="NoSpacing")
    else:
        doc.add_paragraph("\nBody: <none>").runs[0].bold = True

    # Response section
    doc.add_heading("Response", level=2)

    status = response_details.get("Status Code", "")
    doc.add_paragraph(f"\nStatus Code:").runs[0].bold = True
    doc.add_paragraph(str(status), style="NoSpacing")

    doc.add_paragraph("\nResponse Headers:").runs[0].bold = True
    resp_headers = response_details.get("Headers", {})
    for k, v in resp_headers.items():
        doc.add_paragraph(f"{k}: {v}", style="NoSpacing")

    resp_body = response_details.get("Body")
    if resp_body is not None:
        # If body is dict/list try pretty JSON, otherwise string
        doc.add_paragraph("\nBody:").runs[0].bold = True
        if isinstance(resp_body, (dict, list)):
            doc.add_paragraph(_pretty_json(resp_body), style="NoSpacing")
        else:
            # try to parse JSON string
            try:
                parsed = json.loads(resp_body)
                doc.add_paragraph(_pretty_json(parsed), style="NoSpacing")
            except Exception:
                doc.add_paragraph(str(resp_body), style="NoSpacing")
    else:
        doc.add_paragraph("Body: <none>")

    # ---- Assertion Result ----
    doc.add_heading("Assertion Result", level=2)
    doc.add_paragraph(assertion_result, style="NoSpacing")

    # Save document per-test (keeps results separate)
    safe_name = _safe_filename(f"{tc_no}_{tc_name}")
    out_path = os.path.join(output_dir, f"{safe_name}.docx")
    doc.save(out_path)
    return out_path
