import datetime
from docx import Document
import os
import csv

TEST_DATA_PATH = os.path.join(os.path.dirname(__file__), "data", "test_data.csv")
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "utils", "docx_generator", "template.docx")
TEST_RESULTS_DIR = os.path.join(os.path.dirname(__file__), "results")

os.makedirs(TEST_RESULTS_DIR, exist_ok=True)

import re

def safe_filename(name: str) -> str:
    """ Replace unsafe characters in file names with underscores. """
    name = re.sub(r"[^\w\s\-(),.&]", "_", name)  # allow (), space, dash, underscore
    return name.strip()

MAX_PATH_LEN = 257

def safe_path(output_dir: str, tc_no: str, tc_name: str) -> str:
    base_name = f"{safe_filename(tc_no)} {safe_filename(tc_name)}.docx"
    out_path = os.path.join(output_dir, base_name)

    if len(out_path) > MAX_PATH_LEN:
        ext = ".docx"
        suffix = ".." + ext
        max_name_len = MAX_PATH_LEN - len(output_dir) - 1 - len(suffix)

        trimmed_name = f"{safe_filename(tc_no)} {safe_filename(tc_name)}"
        trimmed_name = trimmed_name[:max_name_len]  # trim to fit
        base_name = trimmed_name + suffix
        out_path = os.path.join(output_dir, base_name)

    return out_path

test_cases = []
with open(TEST_DATA_PATH, newline='', encoding='utf-8') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        tc = {
            "no": row["tc_no"],
            "name": row["tc_name"],
            "description": row["tc_description"],
            "execution_date": row.get("execution_date", str(datetime.date.today()))
        }
        test_cases.append(tc)

for tc in test_cases:
    document = Document(TEMPLATE_PATH)
    table = document.tables[0]  # Use the first table on the first page

    # Fill in the table cells
    table.cell(0, 1).text = str(tc["no"])
    table.cell(1, 1).text = tc["name"]

    # Handle multi-line description
    # Clear all existing paragraphs from the cell
    desc_cell = table.cell(2, 1)
    for p in desc_cell.paragraphs:
        p._element.getparent().remove(p._element)

    # Now add new lines
    desc_text = tc["description"].replace("\\n", "\n")  # convert CSV "\n" into real newlines
    for line in desc_text.splitlines():
        desc_cell.add_paragraph(line)

    table.cell(3, 1).text = str(tc["execution_date"])

    out_path = safe_path(TEST_RESULTS_DIR, tc["no"], tc["name"])

    document.save(out_path)
    print(f"✅ Exported test case {tc['no']} to {out_path}")
